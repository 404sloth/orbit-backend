"""
Executive Document Generation Tool.
Creates premium, well-formatted PDFs for Meeting Summaries, RFPs, and SOWs.
Uses fpdf2 for high-quality executive layouts.
"""
import os
import uuid
import datetime
from fpdf import FPDF
from langchain_core.tools import tool
from core.schemas import GenerateReportSchema
from core.logger import logger
from core.config import settings
from core.session import REPORTS_TEMP_DIR








class ExecutivePDF(FPDF):
    def __init__(self):
        super().__init__()
        self.set_margins(20, 20, 20)
        self.set_auto_page_break(auto=True, margin=20)

    def header(self):
        if self.page_no() == 1:
            # Very subtle header on first page
            self.set_font("helvetica", "B", 8)
            self.set_text_color(203, 213, 225) # Dim Slate 300
            self.set_y(10)
            self.cell(0, 10, "ORBIT INTELLIGENCE ENGINE  |  CONFIDENTIAL", align="R")
            self.ln(15)
        

    def footer(self):
        self.set_y(-15)
        self.set_font("helvetica", "I", 8)
        self.set_text_color(226, 232, 240) # Even dimmer Slate 200
        date_str = datetime.datetime.now().strftime("%d %b %Y %H:%M")
        self.cell(0, 10, f"Page {self.page_no()}/{{nb}}  |  Generated: {date_str}", align="C")

def parse_markdown_table(markdown_block):
    """Simple parser to convert markdown table to list of lists."""
    lines = markdown_block.strip().split("\n")
    table_data = []
    for line in lines:
        if "|" in line and "---" not in line:
            # Handle outer pipes correctly
            line = line.strip()
            if line.startswith("|"): line = line[1:]
            if line.endswith("|"): line = line[:-1]
            cells = [c.strip() for c in line.split("|")]
            if cells: table_data.append(cells)
    return table_data

def generate_pdf(doc_type, title, subtitle, content_markdown, filepath):
    """Advanced Markdown-aware PDF generator."""
    pdf = ExecutivePDF()
    pdf.alias_nb_pages()
    pdf.add_page()
    
    # Apply subtle premium background tint
    pdf.set_fill_color(252, 252, 253) # Very light slate/white
    pdf.rect(0, 0, 210, 297, "F")
    
    # 1. Title Section (Professional sizes)
    pdf.set_font("helvetica", "B", 20)
    pdf.set_text_color(15, 23, 42) # Slate 900
    pdf.multi_cell(0, 10, title.upper(), align="L")
    
    if subtitle:
        pdf.set_font("helvetica", "B", 12)
        pdf.set_text_color(197, 165, 114) # Gold
        pdf.cell(0, 8, subtitle, ln=1, align="L")
        
    pdf.ln(2)
    pdf.set_draw_color(197, 165, 114)
    pdf.set_line_width(0.5)
    pdf.line(25, pdf.get_y(), 50, pdf.get_y())
    pdf.ln(8)
    
    # Pre-process markdown: Replace backticks with subtle bold for better PDF look
    import re
    content_markdown = re.sub(r'`([^`]+)`', r'**\1**', content_markdown)
    
    # 2. Content Processing
    # We split by double newlines to handle blocks (text vs tables)
    blocks = content_markdown.split("\n\n")
    
    for block in blocks:
        block = block.strip()
        if not block: continue
        
        # A. Table Detection
        if "|" in block and "---" in block:
            table_data = parse_markdown_table(block)
            if table_data:
                pdf.set_font("helvetica", "", 8) # Smaller font for tables to prevent overflow
                pdf.set_text_color(51, 65, 85)
                
                # Calculate columns to avoid "Not enough horizontal space" errors
                # fpdf2's table() is powerful but sensitive to width
                try:
                    with pdf.table(
                        borders_layout="HORIZONTAL_LINES",
                        cell_fill_color=(248, 250, 252),
                        cell_fill_mode="ROWS",
                        line_height=6,
                        text_align="LEFT",
                        width=pdf.epw
                    ) as table:
                        for row in table_data:
                            row_cells = table.row()
                            for cell_text in row:
                                # Clean cell text to prevent formatting issues
                                clean_text = str(cell_text).strip()
                                row_cells.cell(clean_text)
                    pdf.ln(5)
                except Exception as table_err:
                    logger.warning(f"Table rendering fallback: {table_err}")
                    # Fallback to simple text if table fails
                    pdf.set_font("helvetica", "I", 8)
                    pdf.multi_cell(0, 5, "[Table omitted due to layout constraints - data remains in structured history]")
                    pdf.ln(5)
                continue

        # B. Header Detection
        if block.startswith("#"):
            lines = block.split("\n")
            for line in lines:
                if line.startswith("###"):
                    pdf.set_font("helvetica", "B", 12)
                    pdf.set_text_color(30, 41, 59)
                    pdf.ln(2)
                    pdf.multi_cell(0, 8, line.replace("###", "").strip())
                elif line.startswith("##"):
                    pdf.set_font("helvetica", "B", 14)
                    pdf.set_text_color(15, 23, 42)
                    pdf.ln(4)
                    pdf.multi_cell(0, 10, line.replace("##", "").strip())
                    pdf.ln(1)
                elif line.startswith("#"):
                    pdf.set_font("helvetica", "B", 16)
                    pdf.set_text_color(15, 23, 42)
                    pdf.ln(6)
                    pdf.multi_cell(0, 12, line.replace("#", "").strip())
            continue

        # C. List / Bullet Detection
        if block.startswith("- ") or block.startswith("* "):
            pdf.set_font("helvetica", "", 10)
            pdf.set_text_color(71, 85, 105) # Slate 600
            # Use fpdf2 markdown support for bold/italic in bullets
            try:
                pdf.multi_cell(0, 6, block, markdown=True)
            except Exception:
                pdf.multi_cell(0, 6, block, markdown=False)
            pdf.ln(2)
            continue

        # D. Standard Paragraph
        pdf.set_font("helvetica", "", 10)
        pdf.set_text_color(71, 85, 105)
        # Handle backticks by replacing with subtle bolding or keeping as is if fpdf doesn't support code blocks
        # We use fpdf2's markdown support for **bold**, *italic*, and --strikethrough--
        try:
            pdf.multi_cell(0, 6, block, markdown=True)
        except Exception:
            pdf.multi_cell(0, 6, block, markdown=False)
        pdf.ln(2)
            
    pdf.output(filepath)

def generate_docx(doc_type, title, subtitle, content_markdown, filepath):
    """Helper to generate professional DOCX."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # Title
    t = doc.add_heading(title, 0)
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    if subtitle:
        s = doc.add_paragraph(subtitle)
        s.runs[0].font.color.rgb = RGBColor(197, 165, 114)
        s.runs[0].font.size = Pt(14)
        s.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
    doc.add_paragraph("_" * 60)
    
    lines = content_markdown.split("\n")
    for line in lines:
        line = line.strip()
        if not line: continue
        
        if line.startswith("#"):
            level = line.count("#")
            clean_text = line.replace("#", "").strip()
            h = doc.add_heading(clean_text, level=level)
        elif line.startswith("-") or line.startswith("*"):
            doc.add_paragraph(line[1:].strip(), style='List Bullet')
        else:
            doc.add_paragraph(line)
            
    doc.save(filepath)

def generate_excel(content_markdown, filepath):
    """Helper to generate data Excel."""
    import pandas as pd
    import io
    
    # Try to extract tables from markdown if present, otherwise save text
    lines = content_markdown.split("\n")
    data = []
    for line in lines:
        if "|" in line and "---" not in line:
            parts = [p.strip() for p in line.split("|") if p.strip()]
            if parts: data.append(parts)
            
    if data:
        # Check if we have consistent column counts
        if len(data) > 1:
            try:
                df = pd.DataFrame(data[1:], columns=data[0])
                df.to_excel(filepath, index=False)
                return
            except: pass
            
    # Fallback to simple text dump if no table found or parsing failed
    df = pd.DataFrame({"Content": content_markdown.split("\n")})
    df.to_excel(filepath, index=False)

@tool(args_schema=GenerateReportSchema)
def generate_executive_report(doc_type: str, title: str, content_markdown: str, subtitle: str = None, format: str = "PDF") -> str:
    """
    Generates a premium executive report in the specified format (PDF, DOCX, or EXCEL).
    """
    try:
        output_dir = REPORTS_TEMP_DIR
        os.makedirs(output_dir, exist_ok=True)
        
        safe_title = "".join(x for x in title if x.isalnum() or x in " -_").strip().replace(" ", "_")
        ext = format.lower()
        if ext == "excel": ext = "xlsx"
        
        filename = f"{safe_title}_{uuid.uuid4().hex[:6]}.{ext}"
        filepath = os.path.join(output_dir, filename)
        
        if format.upper() == "PDF":
            generate_pdf(doc_type, title, subtitle, content_markdown, filepath)
        elif format.upper() == "DOCX":
            generate_docx(doc_type, title, subtitle, content_markdown, filepath)
        elif format.upper() == "EXCEL":
            generate_excel(content_markdown, filepath)
        else:
            return f"Error: Unsupported format '{format}'"
            
        download_url = f"http://localhost:8000/reports/download/{filename}"
        logger.info("Professional document generated", format=format, url=download_url)
        
        return f"A professional {format} {doc_type} has been generated: **[{title}]({download_url})**. It is now available in your Executive Artifacts panel."

    except Exception as e:
        logger.error("Document Generation Failed", error=str(e))
        return f"Error generating document: {str(e)}"
